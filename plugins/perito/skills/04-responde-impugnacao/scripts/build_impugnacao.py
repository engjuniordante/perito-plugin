# -*- coding: utf-8 -*-
"""
build_impugnacao.py — monta os esclarecimentos/impugnação (.docx) a partir de um JSON.
Uso: python3 build_impugnacao.py <template-impugnacao.docx> <data.json> <saida.docx>

O MODELO produz só o JSON (campos + corpo dos esclarecimentos, vindos do output do NLM).
O SCRIPT monta o .docx (escalares + dropdown de parte + corpo formatado), determinístico.

JSON:
{
  "perito_nome": "Irineu de Freitas Branco Junior",
  "scalars": { "CIDADE_VARA":..., "NUMERO_PROCESSO":..., "NOME_RECLAMANTE":...,
               "NOME_RECLAMADA":..., "ID_IMPUGNACAO":..., "DATA_EXTENSO":... },
  "parte_impugnante": "Reclamada",   // ou "Reclamante" (dropdown SDT)
  "esclarecimentos": [               // corpo; o script auto-formata:
     "ESCLARECIMENTOS SOLICITADOS PELA RECLAMADA",   // -> título (negrito)
     "[parágrafo de fundamentação, se houver]",
     "1- [texto do quesito]",
     "Resposta: [texto da resposta]",                // -> "Resposta:" em negrito
     ...
  ]
}
"""
import sys, json, re, os, ntpath
from copy import deepcopy

# Piso 3.9 (anotações builtin) + stdout/err UTF-8: no Windows a saída capturada cai em
# cp1252 (Python <3.15) e um emoji do relatório mataria o script com UnicodeEncodeError.
if sys.version_info < (3, 9):
    sys.exit('Python 3.9+ é necessário (este ambiente tem %d.%d).' % sys.version_info[:2])
for _s in (sys.stdout, sys.stderr):
    if _s is not None and hasattr(_s, 'reconfigure'):
        _s.reconfigure(encoding='utf-8', errors='replace')

# --- auto-provisionamento de dependências (sandbox efêmero do Cowork) ---
def _ensure(pkgs):
    import importlib, importlib.util, subprocess, sys as _sys
    falta = [pip for mod, pip in pkgs if importlib.util.find_spec(mod) is None]
    if not falta:
        return
    cmd = [_sys.executable, '-m', 'pip', 'install', *falta]
    # PEP 668 (Linux E macOS/Homebrew): pip recusa fora de venv → repete com a flag
    if subprocess.run(cmd, check=False).returncode != 0:
        subprocess.run(cmd + ['--break-system-packages'], check=False)
    importlib.invalidate_caches()
    resta = [pip for mod, pip in pkgs if importlib.util.find_spec(mod) is None]
    if resta:
        _sys.exit('dependência ausente: %s — instale com:\n  %s -m pip install %s'
                  % (', '.join(resta), _sys.executable, ' '.join(resta)))
_ensure([('docx', 'python-docx')])

import docx
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

def all_paragraphs(document):
    out = []
    def rec(tbls):                     # recursivo: pega marcador em tabela aninhada (paridade c/ skills 02/03)
        for t in tbls:
            for row in t.rows:
                for cell in row.cells:
                    out.extend(cell.paragraphs); rec(cell.tables)
    out.extend(document.paragraphs); rec(document.tables)
    for sec in document.sections:
        out.extend(sec.header.paragraphs); out.extend(sec.footer.paragraphs)
    return out

def replace_scalar(document, mapping):
    for p in all_paragraphs(document):
        for run in p.runs:
            if '{{' in run.text:
                t = run.text
                for k, v in mapping.items():
                    if k in t: t = t.replace(k, str(v))
                run.text = t

def set_sdt_party(document, parte):
    """Define o valor exibido do dropdown SDT (alias 'Partes')."""
    body = document.element.body
    for sdt in body.iter(qn('w:sdt')):
        pr = sdt.find(qn('w:sdtPr'))
        alias = pr.find(qn('w:alias')) if pr is not None else None
        if alias is not None and alias.get(qn('w:val')) == 'Partes':
            content = sdt.find(qn('w:sdtContent'))
            ts = list(content.iter(qn('w:t')))
            if ts:
                ts[0].text = parte
                for extra in ts[1:]:
                    extra.text = ''
            return True
    return False

def render_esclarecimentos(document, itens):
    # localizar o parágrafo do marcador
    target = None
    for p in document.paragraphs:
        if '{{ESCLARECIMENTOS_CORPO}}' in p.text:
            target = p; break
    if target is None:
        return False
    # rPr base (do marcador)
    rpr = None
    if target.runs:
        el = target.runs[0]._r.find(qn('w:rPr'))
        if el is not None: rpr = deepcopy(el)
    base_p = target._p

    def new_par_after(anchor_p):
        np = deepcopy(base_p)
        for r in np.findall(qn('w:r')): np.remove(r)
        anchor_p.addnext(np)
        return np

    def add_run(par_el, text, bold=False):
        r = docx.oxml.OxmlElement('w:r')
        rp = deepcopy(rpr) if rpr is not None else docx.oxml.OxmlElement('w:rPr')
        # limpar bold anterior
        for b in rp.findall(qn('w:b')): rp.remove(b)
        if bold:
            rp.insert(0, docx.oxml.OxmlElement('w:b'))
        r.append(rp)
        t = docx.oxml.OxmlElement('w:t'); t.set(qn('xml:space'), 'preserve'); t.text = text
        r.append(t)
        par_el.append(r)

    anchor = base_p
    first = True
    for item in itens:
        par_el = base_p if first else new_par_after(anchor)
        if not first:
            anchor = par_el
        else:
            for r in par_el.findall(qn('w:r')): par_el.remove(r)  # limpar marcador
            first = False
            anchor = par_el
        s = item.strip()
        if s.upper().startswith('ESCLARECIMENTOS SOLICITADOS'):
            add_run(par_el, s.upper(), bold=True)
        elif s.lower().startswith('resposta:'):
            add_run(par_el, 'Resposta:', bold=True)
            rest = s[len('resposta:'):]
            add_run(par_el, rest if rest.startswith(' ') else ' ' + rest.lstrip(), bold=False)
        else:
            add_run(par_el, s, bold=False)
    return True

def _resolve_template(template_path):
    """Fallback BUNDLED — no Cowork o bash NÃO enxerga o Drive; cai no template do plugin
    (assets/templates/). Nativo (bash vê o Drive) usa o caminho vivo."""
    if template_path and os.path.isfile(template_path):
        return template_path
    here = os.path.dirname(os.path.abspath(__file__))
    tdir = os.path.join(here, '..', 'assets', 'templates')
    # ntpath.basename entende tanto '/' (POSIX) quanto '\' (Windows/Drive);
    # os.path.basename no Linux ignora '\' e devolveria o caminho inteiro do Drive.
    cand = os.path.join(tdir, ntpath.basename(template_path or ''))
    if os.path.isfile(cand):
        print('ℹ️  template do Drive inacessível (bash do Cowork) — usando o BUNDLED: %s'
              % ntpath.basename(cand))
        return cand
    raise SystemExit('template não encontrado: nem "%s" nem bundled em %s' % (template_path, tdir))


def build(template, data_path, out_path):
    data = json.load(open(data_path, encoding='utf-8'))
    doc = docx.Document(_resolve_template(template))
    warnings = []
    fatal = []   # compromete o documento → bloqueia a saída (exit 2)

    scal = {(k if k.startswith('{{') else '{{%s}}' % k): v
            for k, v in data.get('scalars', {}).items()}
    replace_scalar(doc, scal)

    parte = data.get('parte_impugnante', 'Reclamada')
    if not set_sdt_party(doc, parte):
        warnings.append('dropdown SDT "Partes" não encontrado — ajuste manualmente')

    if not render_esclarecimentos(doc, data.get('esclarecimentos', [])):
        fatal.append('marcador {{ESCLARECIMENTOS_CORPO}} não encontrado — corpo dos esclarecimentos não entrou')

    full = '\n'.join(p.text for p in all_paragraphs(doc))
    residual = sorted(set(re.findall(r'\{\{[^}]+\}\}', full)))
    if residual: fatal.append('MARCADORES RESIDUAIS: ' + ', '.join(residual))
    perito = data.get('perito_nome', 'Irineu de Freitas Branco Junior')
    if perito not in full: fatal.append('IDENTIDADE: perito (%s) ausente do documento' % perito)
    for bad in data.get('nomes_proibidos', []):
        if bad in full: fatal.append('VAZAMENTO: "%s" presente no documento' % bad)

    if fatal:
        print('\n❌ NÃO GERADO — problema(s) que comprometem os esclarecimentos:')
        for f in fatal: print('  -', f)
        print('   Nenhum arquivo foi salvo. Corrija e rode de novo.')
        return False

    doc.save(out_path)
    print('OK ->', out_path, '| parte impugnante:', parte)
    if warnings:
        print('\n⚠ AVISOS (revise antes de assinar):')
        for w in warnings: print('  -', w)
    else:
        print('✅ VALIDAÇÃO OK: sem marcador residual, identidade presente, sem vazamento.')
    return True

if __name__ == '__main__':
    if len(sys.argv) != 4:
        print('uso: python3 build_impugnacao.py <template> <data.json> <saida.docx>'); sys.exit(1)
    sys.exit(0 if build(sys.argv[1], sys.argv[2], sys.argv[3]) else 2)
