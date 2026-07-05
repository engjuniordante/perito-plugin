#!/usr/bin/env python3
"""Regressão do redator de impugnação (build_impugnacao.py).
Trava os fixes v1.0.66: ntpath.basename (path do Drive Windows) e all_paragraphs recursivo (bug 2)."""
import os, sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent))
import docx
import build_impugnacao as bi

FALHAS = []


def check(cond, msg):
    print(("  ✓ " if cond else "  ✗ FALHOU: ") + msg)
    if not cond:
        FALHAS.append(msg)


def _bundled_dir():
    return os.path.join(os.path.dirname(os.path.abspath(bi.__file__)), '..', 'assets', 'templates')


def t_resolve_windows():
    print("T1 — _resolve_template com path do Drive Windows cai no BUNDLED")
    win = r'G:\Meu Drive\pericias\template-impugnacao.docx'
    r = bi._resolve_template(win)
    check(os.path.isfile(r), 'resolveu para arquivo existente')
    check(os.path.basename(r) == 'template-impugnacao.docx', 'basename correto (ntpath isolou o nome)')

    try:
        bi._resolve_template(r'G:\x\template-que-nao-existe.docx')
        check(False, 'basename inexistente deveria abortar (SystemExit)')
    except SystemExit:
        check(True, 'basename inexistente aborta — nao substitui por outro template')


def t_all_paragraphs_recursivo():
    print("T2 — all_paragraphs enxerga marcador em tabela ANINHADA (bug 2)")
    d = docx.Document()
    outer = d.add_table(rows=1, cols=1)
    inner = outer.cell(0, 0).add_table(rows=1, cols=1)
    inner.cell(0, 0).paragraphs[0].add_run('{{MARCADOR_ANINHADO}}')
    textos = [p.text for p in bi.all_paragraphs(d)]
    check(any('{{MARCADOR_ANINHADO}}' in t for t in textos),
          'marcador dentro de tabela-em-tabela aparece (recursao restaurada)')


if __name__ == '__main__':
    t_resolve_windows(); t_all_paragraphs_recursivo()
    print()
    if FALHAS:
        print('FALHOU (%d): %s' % (len(FALHAS), '; '.join(FALHAS))); sys.exit(1)
    print('OK — todos os testes do build_impugnacao passaram'); sys.exit(0)
