#!/usr/bin/env python3
"""Regressão do redator insal/peric (build_laudo.py).
Trava os fixes v1.0.66: ntpath.basename (path do Drive Windows) e str() no replace_scalar."""
import os, sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent))
import docx
import build_laudo as bl

FALHAS = []


def check(cond, msg):
    print(("  ✓ " if cond else "  ✗ FALHOU: ") + msg)
    if not cond:
        FALHAS.append(msg)


def t_resolve_windows():
    print("T1 — _resolve_template com path do Drive Windows cai no BUNDLED")
    win = r'G:\Meu Drive\pericias\template-insalubridade.docx'
    r = bl._resolve_template(win)
    check(os.path.isfile(r), 'resolveu para arquivo existente')
    check(os.path.basename(r) == 'template-insalubridade.docx', 'basename correto (ntpath isolou o nome)')

    posix = '/caminho/inexistente/template-periculosidade.docx'
    r2 = bl._resolve_template(posix)
    check(os.path.basename(r2) == 'template-periculosidade.docx', 'path POSIX inexistente tambem cai no bundled')

    try:
        bl._resolve_template(r'G:\x\template-que-nao-existe.docx')
        check(False, 'basename inexistente deveria abortar (SystemExit)')
    except SystemExit:
        check(True, 'basename inexistente aborta — nao substitui por outro template')


def t_gate_tipo_windows():
    print("T2 — _gate_tipo aceita basename de path Windows (fix da linha 107)")
    doc = docx.Document(os.path.join(bl._BUNDLED_TEMPLATES, 'template-insalubridade.docx'))
    win = r'G:\Meu Drive\pericias\template-insalubridade.docx'
    errs, _tipo, _intent = bl._gate_tipo(doc, win, 'insalubridade')
    tipo_errs = [e for e in errs if 'TEMPLATE PEDIDO' in e]
    check(not tipo_errs, 'sem falso "TIPO x TEMPLATE PEDIDO" (ntpath.basename == template esperado)')


def t_replace_scalar_nao_str():
    print("T3 — replace_scalar coage valor nao-str sem crash (bug 1)")
    d = docx.Document()
    p = d.add_paragraph('{{EPI_ANO_1}}')
    try:
        bl.replace_scalar(d, {'{{EPI_ANO_1}}': 2021})  # int, como o JSON pode emitir
        crashed = False
    except TypeError:
        crashed = True
    check(not crashed, 'nao crasha com int no mapping')
    check('2021' in p.text and '{{EPI_ANO_1}}' not in p.text, 'substituiu pelo str do numero')


FORM = """## ▶ EPIs FORNECIDOS

### NR-6 — Comprovação
- Ficha de EPI — registro do fornecimento 🔄 — [X] Sim  [ ] Não
- Anotação do respectivo C.A. 🔄 — [ ] Sim  [X] Não
- Treinamento e orientação 🔄 — [X] Sim  [ ] Não
- Frequência regular de fornecimento 🔄 — [X] Sim  [ ] Não
- Adequado ao risco ambiental 👤 (perito) — [ ] Sim  [ ] Não
- Fiscalização do uso 👤 (perito) — [ ] Sim  [ ] Não

## ▶ AGENTES — INSALUBRIDADE (NR-15)

### A. RUÍDO (Anexos 1 e 2)
- Status: [ ] Ausente  [X] Presente
- Nível medido (dB): 97,0
- Neutralizado pelo EPI: [X] Sim  [ ] Não

### H. FRIO (Anexo 9)
- Status: [X] Ausente  [ ] Presente
"""


def t_parsers_formulario():
    print("T4 — parsing do formulário (checkbox do perito)")
    check(bl.neutralizacao_por_agente(FORM) == {'ANALISE_RUIDO_CONTINUO': 'SIM',
                                                'ANALISE_RUIDO_IMPACTO': 'SIM'},
          'lê "Neutralizado: [X] Sim" do bloco A (os 2 marcadores de ruído)')
    check(bl.nr6_do_form(FORM) == {'ficha': 'SIM', 'ca': 'NAO', 'treinamento': 'SIM',
                                   'frequencia': 'SIM'},
          'NR-6: 4 linhas marcadas; as 2 do perito (branco) ficam fora')
    novo = FORM.replace('- Neutralizado pelo EPI: [X] Sim  [ ] Não',
                        '- Neutralizado pelo EPI durante TODO o imprescrito: [X] Sim  [ ] Não (nunca, ou só em parte)')
    check(bl.neutralizacao_por_agente(novo) == {'ANALISE_RUIDO_CONTINUO': 'SIM', 'ANALISE_RUIDO_IMPACTO': 'SIM'},
          'rótulo desambiguado (com sufixo) → mesmo resultado')
    parcial = FORM.replace('- Neutralizado pelo EPI: [X] Sim  [ ] Não',
                           '- Neutralizado pelo EPI durante TODO o imprescrito: [ ] Sim  [X] Não (nunca, ou só em parte)')
    check(bl.neutralizacao_por_agente(parcial) == {'ANALISE_RUIDO_CONTINUO': 'NAO', 'ANALISE_RUIDO_IMPACTO': 'NAO'},
          'cobertura parcial → NAO (o "Não" do rótulo não confunde o checkbox)')
    check(bl.gate_formulario({'blocks': {'ANALISE_RUIDO_CONTINUO': ['Caracterizada a insalubridade em grau médio '
                                                                   'de 01/2024 a 03/2024.']}, 'nr6': {}}, parcial) == [],
          'PARCIAL + conclusão insalubre no recorte → gate NÃO bloqueia (ressalva do Codex)')
    check(bl.caracteriza_insalubridade(['Atividades insalubres em grau médio.']), 'caracteriza: positivo')
    check(bl.caracteriza_insalubridade(['Caracterizada a insalubridade no período.']), 'caracteriza: "Caracterizada a insalubridade"')
    check(not bl.caracteriza_insalubridade(['Descaracterizada a insalubridade.']), 'não caracteriza: "Descaracterizada"')
    check(not bl.caracteriza_insalubridade(['Não foi caracterizada a insalubridade.']), 'não caracteriza: "Não foi caracterizada"')


def t_gate_formulario():
    print("T5 — gates 1.7/1.8: JSON não contraria checkbox do perito (bug do VICTOR)")
    base = {'blocks': {}, 'nr6': {'ficha': 'SIM', 'ca': 'NAO', 'treinamento': 'SIM',
                                  'frequencia': 'SIM', 'adequado': '', 'fiscalizacao': ''}}

    ok_data = dict(base, blocks={'{{ANALISE_RUIDO_CONTINUO}}': ['Descaracterizada a insalubridade.']})
    check(bl.gate_formulario(ok_data, FORM) == [], 'JSON fiel ao formulário → passa')

    # 1.7 — a inversão do run real
    inv = dict(base, blocks={'ANALISE_RUIDO_CONTINUO': [
        'Não houve treinamento comprovado.', 'Caracterizada a insalubridade em grau médio.']})
    errs = bl.gate_formulario(inv, FORM)
    check(len(errs) == 1 and 'ANALISE_RUIDO_CONTINUO' in errs[0], 'form "Neutralizado: Sim" + bloco caracteriza → erro')

    # 1.8 — as duas células que o modelo virou
    virado = dict(ok_data, nr6=dict(base['nr6'], treinamento='NAO', ca='SIM'))
    errs = bl.gate_formulario(virado, FORM)
    check(len(errs) == 2 and all('nr6[' in e for e in errs), 'JSON vira treinamento e C.A. contra o form → 2 erros')

    # linha que o perito deixou em branco: o modelo pode preencher
    branco = dict(ok_data, nr6=dict(base['nr6'], fiscalizacao='NAO'))
    check(bl.gate_formulario(branco, FORM) == [], 'linha em branco no form preenchida pelo modelo → não bloqueia')


def t_neutraliza_genero_conclusao():
    """Ementa/conclusao sempre na forma neutra 'O(A) Reclamante ... exposto(a)' — o modelo
    flexiona por conta propria e o laudo saia misturado com a negativa (que ja nasce neutra)."""
    print("T — forma neutra na ementa/conclusao")

    real = ("A Reclamante ficava exposta e mantinha contato com agentes biologicos no periodo "
            "em que exerceu a funcao de Controladora de Acesso na Portaria, sendo caracterizada "
            "a insalubridade em grau maximo, correspondente ao percentual de 40%.")
    out, n = bl._neutraliza_genero_conclusao([real])
    check(n == 1 and out[0].startswith('O(A) Reclamante ficava exposto(a)'),
          'feminino flexionado vira "O(A) Reclamante ficava exposto(a)": %r' % out[0][:60])
    check('Controladora de Acesso' in out[0],
          'nome da funcao (cargo do processo) NAO e mexido')

    out, n = bl._neutraliza_genero_conclusao(
        ["O Reclamante ficava exposto a ruido, sendo caracterizada a insalubridade."])
    check(n == 1 and out[0].startswith('O(A) Reclamante ficava exposto(a)'),
          'masculino flexionado tambem e neutralizado')

    ja_neutro = "O(A) Reclamante ficava exposto(a) a agentes quimicos, sendo caracterizada."
    out, n = bl._neutraliza_genero_conclusao([ja_neutro])
    check(n == 0 and out[0] == ja_neutro,
          'item ja neutro fica intacto (sem "exposto(a)(a)"): %r' % out[0][:60])

    # idempotencia: normalizar o proprio resultado nao muda mais nada
    uma, _ = bl._neutraliza_genero_conclusao(["A Reclamante ficava exposta a ruido."])
    duas, n2 = bl._neutraliza_genero_conclusao(uma)
    check(n2 == 0 and duas == uma, 'segunda passada e no-op (idempotente)')

    out, _ = bl._neutraliza_genero_conclusao(
        ["A Reclamada confirmou as atividades da Reclamante em area exposta a intemperies."])
    check('A Reclamada confirmou' in out[0], '"Reclamada" (a empresa) nao e neutralizada')
    check('area exposta a intemperies' in out[0],
          'participio solto ("area exposta") preservado — so concorda com Reclamante')
    check('do(a) Reclamante' in out[0], '"da Reclamante" vira "do(a) Reclamante"')

    out, _ = bl._neutraliza_genero_conclusao(
        ["Nao foi constatada, nas atividades exercidas pela Reclamante, exposicao ao calor."])
    check('pelo(a) Reclamante' in out[0], '"pela Reclamante" vira "pelo(a) Reclamante"')

    # a negativa que o script acrescenta ja nasce neutra — tem de passar incolume
    out, n = bl._neutraliza_genero_conclusao([bl._CONCL_NEG_PERIC, bl._CONCL_NEG_INSAL])
    check(n == 0 and out == [bl._CONCL_NEG_PERIC, bl._CONCL_NEG_INSAL],
          'frases-padrao de negativa do proprio script ficam intactas')


if __name__ == '__main__':
    t_resolve_windows(); t_gate_tipo_windows(); t_replace_scalar_nao_str()
    t_parsers_formulario(); t_gate_formulario(); t_neutraliza_genero_conclusao()
    print()
    if FALHAS:
        print('FALHOU (%d): %s' % (len(FALHAS), '; '.join(FALHAS))); sys.exit(1)
    print('OK — todos os testes do build_laudo passaram'); sys.exit(0)
