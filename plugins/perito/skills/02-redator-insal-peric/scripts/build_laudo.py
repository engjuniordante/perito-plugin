# -*- coding: utf-8 -*-
"""
build_laudo.py — monta o laudo .docx do Irineu a partir de um JSON de conteúdo.
Uso:  python3 build_laudo.py <template.docx> <laudo-data.json> <saida.docx>

O MODELO (skill) produz só o JSON (dados do formulário + análises).
Este script faz toda a manipulação do .docx — determinístico, sem regex no LLM.

Estrutura do JSON:
{
  "perito_nome": "Irineu de Freitas Branco Junior",   // checagem de identidade
  "scalars":  { "VARA":..., "PROCESSO":..., "RECLAMANTE":..., ... },
  "blocks":   { "LISTA_PARTICIPANTES":[...], "ATIVIDADES_POR_FUNCAO":[...],
                "ANALISE_*":[...], "CONCLUSAO_ITENS":[...], "QUESITOS_RECLAMANTE":[...], ... },
  "identificacao": [ [Funcao,Setor,Inicio,Termino,Autuacao,ImprInicio,ImprTermino], ... ],
  "epi": { "anos": ["2023","2024","2025"], "linhas": [ [Desc,Agente,CA,v1,v2,v3], ... ] },
  "nr6": { "ficha":"SIM","ca":"SIM","treinamento":"SIM","adequado":"","frequencia":"NAO","fiscalizacao":"" },
  "vibracao": [ ["Trator de Transbordo","1,00","16,50"], ["Colhedora","0,60","7,90"] ]   // opcional
}
"""
import sys, json, re, os, ntpath, sqlite3, unicodedata
from copy import deepcopy
from pathlib import Path

# Piso 3.9 (anotações builtin) + stdout/err UTF-8: no Windows a saída capturada cai em
# cp1252 (Python <3.15) e um emoji do relatório mataria o script com UnicodeEncodeError.
if sys.version_info < (3, 9):
    sys.exit('Python 3.9+ é necessário (este ambiente tem %d.%d).' % sys.version_info[:2])
for _s in (sys.stdout, sys.stderr):
    if _s is not None and hasattr(_s, 'reconfigure'):
        _s.reconfigure(encoding='utf-8', errors='replace')

# --- auto-provisionamento de dependências (sandbox efêmero do Cowork) ---
def _ensure(pkgs):
    import importlib, importlib.util, subprocess, sys as _sys
    falta = [pip for mod, pip in pkgs if importlib.util.find_spec(mod) is None]
    if not falta:
        return
    cmd = [_sys.executable, '-m', 'pip', 'install', *falta]
    # PEP 668 (Linux E macOS/Homebrew): pip recusa fora de venv → repete com a flag
    if subprocess.run(cmd, check=False).returncode != 0:
        subprocess.run(cmd + ['--break-system-packages'], check=False)
    importlib.invalidate_caches()
    resta = [pip for mod, pip in pkgs if importlib.util.find_spec(mod) is None]
    if resta:
        _sys.exit('dependência ausente: %s — instale com:\n  %s -m pip install %s'
                  % (', '.join(resta), _sys.executable, ' '.join(resta)))
_ensure([('docx', 'python-docx')])

import docx
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from docx.shared import Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

# --- resolução de caminhos com fallback BUNDLED (Cowork: o bash NÃO enxerga o Drive) ---
# O template .docx e a base EPI podem chegar por caminho do Drive (config). No sandbox do
# Cowork o bash não alcança o Drive → cai na cópia BUNDLED no plugin (mesma tática da base
# EPI da skill 01). Nativo (bash vê o Drive) usa o caminho vivo; Cowork usa o bundled.
_HERE = os.path.dirname(os.path.abspath(__file__))
_BUNDLED_TEMPLATES = os.path.join(_HERE, '..', 'assets', 'templates')
_BUNDLED_EPI = os.path.normpath(os.path.join(_HERE, '..', '..', '01-extrator', 'assets', '04-EPIs'))


def _list_bundled_templates():
    d = _BUNDLED_TEMPLATES
    return [f for f in os.listdir(d) if f.lower().endswith('.docx')] if os.path.isdir(d) else []


def _resolve_template(template_path):
    """(1) caminho dado, se for arquivo legível; (2) bundled assets/templates/<basename>.
    O fallback só aceita o BUNDLED de MESMO basename — nunca troca de tipo em silêncio.
    Se o basename pedido não existe bundled, falha (não substitui por outro template)."""
    if template_path and os.path.isfile(template_path):
        return template_path
    # ntpath.basename entende '/' (POSIX) e '\' (Windows/Drive); os.path.basename no Linux
    # ignora '\' e devolveria o caminho inteiro do Drive → fallback bundled nunca casaria.
    cand = os.path.join(_BUNDLED_TEMPLATES, ntpath.basename(template_path or ''))
    if os.path.isfile(cand):
        print('ℹ️  template do Drive inacessível (bash do Cowork) — usando o BUNDLED: %s'
              % ntpath.basename(cand))
        return cand
    raise SystemExit('template não encontrado: nem "%s" nem bundled em %s\n'
                     '   (bundled disponíveis: %s)'
                     % (template_path, _BUNDLED_TEMPLATES,
                        ', '.join(sorted(_list_bundled_templates())) or 'nenhum'))


# --- gate de tipo: o template resolvido TEM que casar com o tipo do laudo ---
# Blindagem contra o b.o. real (laudo só-insalubridade saiu num template insal+peric,
# com seção 7 de periculosidade fantasma). Cruza duas fontes de verdade com o CONTEÚDO
# real do template resolvido (marcadores {{ANALISE_PERIC_*}} vs NR-15):
#   (a) tipo_laudo declarado no JSON  → vem do "▶ TIPO DE LAUDO" do formulário;
#   (b) o basename do template pedido no comando.
# Qualquer divergência = falha ANTES de gravar. Pega tanto o hardcode do comando quanto
# um fallback BUNDLED que tenha trocado de tipo.
_TIPO_CANON = {
    'insalubridade': 'insalubridade', 'insal': 'insalubridade',
    'periculosidade': 'periculosidade', 'peric': 'periculosidade',
    'insalubridade + periculosidade': 'insal-peric', 'insal-peric': 'insal-peric',
    'insal+peric': 'insal-peric', 'insalubridade e periculosidade': 'insal-peric',
    'insal e peric': 'insal-peric',
}
# tipo -> (basename esperado, precisa NR-15, precisa NR-16)
_TIPO_TEMPLATE = {
    'insalubridade':  ('template-insalubridade.docx',  True,  False),
    'periculosidade': ('template-periculosidade.docx', False, True),
    'insal-peric':    ('template-insal-peric.docx',    True,  True),
}


def _canon_tipo(s):
    return _TIPO_CANON.get(re.sub(r'\s+', ' ', str(s).strip().lower())) if s else None


def _gate_tipo(doc, requested_path, tipo_laudo):
    """Devolve (erros, tipo_canon, intent). Detecta o tipo pelo CONTEÚDO do template
    (marcadores NR-15 vs {{ANALISE_PERIC_*}}) e confronta com o declarado/pedido."""
    raw = '\n'.join(p.text for p in all_paragraphs(doc))
    tmpl_nr15 = bool(re.search(r'\{\{ANALISE_(?!PERIC_)[A-Z]', raw))
    tmpl_nr16 = bool(re.search(r'\{\{ANALISE_PERIC_', raw))
    req_base = ntpath.basename(requested_path or '')  # aceita path do Drive Windows (vide _resolve_template)
    tipo = _canon_tipo(tipo_laudo)
    intent = tipo  # sem tipo_laudo, inferir do basename pedido
    if intent is None:
        for t, (base, _n15, _n16) in _TIPO_TEMPLATE.items():
            if base == req_base:
                intent = t
                break
    errs = []
    if tipo and req_base:  # (1) tipo declarado × basename pedido (pega hardcode do comando)
        exp_base = _TIPO_TEMPLATE[tipo][0]
        if req_base != exp_base:
            errs.append('TIPO × TEMPLATE PEDIDO: laudo é "%s" mas o comando passou "%s" '
                        '(esperado "%s"). Corrija o 1º argumento.' % (tipo, req_base, exp_base))
    if intent:  # (2) intenção × conteúdo REAL do template (pega fallback trocado de tipo)
        _b, need15, need16 = _TIPO_TEMPLATE[intent]
        if need16 != tmpl_nr16:
            errs.append('TEMPLATE ERRADO (periculosidade): tipo "%s" %s seção NR-16, mas o '
                        'template resolvido %s marcadores de periculosidade. '
                        'Provável BUNDLED trocado — confira assets/templates/.'
                        % (intent, 'EXIGE' if need16 else 'NÃO admite',
                           'NÃO tem' if need16 else 'CONTÉM'))
        if need15 != tmpl_nr15:
            errs.append('TEMPLATE ERRADO (insalubridade): tipo "%s" %s seção NR-15, mas o '
                        'template resolvido %s marcadores de NR-15.'
                        % (intent, 'EXIGE' if need15 else 'NÃO admite',
                           'NÃO tem' if need15 else 'CONTÉM'))
    return errs, tipo, intent

# ---------------- descaracterização-padrão por agente AUSENTE ----------------
# O modelo só emite os ANALISE_* dos agentes PRESENTES; este mapa preenche
# automaticamente os AUSENTES com a redação verbatim do Irineu (extraída dos
# laudos dele). Cada valor é a lista de parágrafos do bloco.
# ⚠ 4 entradas (RUIDO_CONTINUO, VIBRACOES, QUIM_QUALITATIVOS, PERIC_INFLAMAVEIS)
# foram derivadas do padrão (são quase sempre PRESENTES, raramente caem aqui).
# ---------------- fronteira FORMULÁRIO × modelo (o checkbox do perito manda) ------
# Portado do squad laudo-pericial (gates 1.7/1.8, 09/07/2026). Origem: um redator leu
# uma regra de outro caso ("assinatura na ficha ≠ treinamento"), virou a célula de
# Treinamento da tabela NR-6 contra o [X] Sim do perito e — pela regra "qualquer NÃO
# estrutural = EPI não elide" — inverteu o veredito de ruído de um agente que o
# formulário marcava como "Neutralizado pelo EPI: [X] Sim".
# O formulário é a autoridade sobre o FATO. Divergiu → não gera; leva ao perito.

# letra do bloco de agente no formulário → marcador ANALISE_* do template
_NEUTRAL_MARCADOR = {
    'A': ['ANALISE_RUIDO_CONTINUO', 'ANALISE_RUIDO_IMPACTO'], 'B': ['ANALISE_CALOR'],
    'D': ['ANALISE_RAD_IONIZANTES'], 'E': ['ANALISE_HIPERBARICAS'],
    'F': ['ANALISE_RAD_NAO_IONIZANTES'], 'G': ['ANALISE_VIBRACOES'],
    'H': ['ANALISE_FRIO'], 'I': ['ANALISE_UMIDADE'],
    'J': ['ANALISE_QUIM_QUANTITATIVOS'], 'K': ['ANALISE_POEIRAS_MINERAIS'],
    'L': ['ANALISE_QUIM_QUALITATIVOS'], 'M': ['ANALISE_AGENTES_BIOLOGICOS'],
}  # C (Iluminação, An.4) é revogado — sem veredicto de neutralização

# rótulo da linha NR-6 no formulário → chave do JSON (`nr6`). Ordem importa:
# "Frequência regular de fornecimento" contém "fornecimento".
_NR6_CHAVES = (('frequencia', ('frequencia',)), ('treinamento', ('treinamento',)),
               ('fiscalizacao', ('fiscaliza',)), ('adequado', ('adequa',)),
               ('ca', ('c.a', 'anotacao')), ('ficha', ('ficha', 'fornecimento')))

_AGENTE_LETRA_RE = re.compile(r'^\s*#{0,4}\s*([A-M])\.\s+\S')
# tolera sufixo no rótulo ("...pelo EPI durante TODO o imprescrito:") — formulários
# antigos trazem só "pelo EPI:"; ambos casam
_NEUTRAL_LABEL_RE = re.compile(r'(?i)^\s*[-*•]?\s*Neutralizado\s+pelo\s+EPI\b[^:]*:')
_SECAO_RE = re.compile(r'^\s*(?:#{1,4}\s*)?▶')


def _sem_acento(s):
    import unicodedata
    return unicodedata.normalize('NFKD', s or '').encode('ASCII', 'ignore').decode('ASCII').lower()


def _checkbox_sim_nao(line):
    """'SIM' | 'NAO' | None conforme o [X] da linha."""
    for chk, opt in re.findall(r'\[\s*([xX ])\s*\]\s*(Sim|N[ãa]o)', line or ''):
        if chk.strip():
            return 'SIM' if opt.lower() == 'sim' else 'NAO'
    return None


def _secao_form(text, titulo):
    """Corpo da seção '▶ TITULO' até a próxima seção ▶ (ou fim)."""
    linhas = (text or '').splitlines()
    h = next((i for i, l in enumerate(linhas)
              if _SECAO_RE.match(l) and titulo.upper() in _sem_acento(l).upper()), None)
    if h is None:
        return ''
    corpo = []
    for l in linhas[h + 1:]:
        if _SECAO_RE.match(l):
            break
        corpo.append(l)
    return '\n'.join(corpo)


def caracteriza_insalubridade(paragrafos):
    """True quando o bloco ANALISE_* afirma insalubridade. Apaga primeiro as formas
    negativas ('descaracterizada', 'não caracterizada', 'não insalubre') para que a
    menção ao radical dentro delas não conte como afirmação."""
    t = _sem_acento(' '.join(paragrafos if isinstance(paragrafos, list) else [str(paragrafos)]))
    t = re.sub(r'descaracterizad[ao]\w*', ' ', t)
    t = re.sub(r'nao\s+(?:se\s+|foi\s+|foram\s+|resta\s+|restou\s+)?caracteriz\w*', ' ', t)
    t = re.sub(r'nao\s+(?:sao\s+|e\s+)?insalubr\w*', ' ', t)
    return bool(re.search(r'caracterizad[ao]\s+a?\s*insalubridade|insalubr\w*\s+em\s+grau', t))


def neutralizacao_por_agente(form_text):
    """{marcador ANALISE_*: 'SIM'|'NAO'} por bloco de agente cuja linha 'Neutralizado
    pelo EPI' traz um [X]. Linha em branco → ausente do dict (perito não opinou)."""
    out, letra = {}, None
    for l in _secao_form(form_text, 'AGENTES').splitlines():
        m = _AGENTE_LETRA_RE.match(l)
        if m:
            letra = m.group(1).upper()
            continue
        if letra and _NEUTRAL_LABEL_RE.match(l):
            v = _checkbox_sim_nao(l)
            if v:
                for mk in _NEUTRAL_MARCADOR.get(letra, []):
                    out[mk] = v
    return out


def _nr6_chave(rotulo):
    s = _sem_acento(rotulo)
    for chave, pistas in _NR6_CHAVES:
        if any(p in s for p in pistas):
            return chave
    return None


def nr6_do_form(form_text):
    """{chave: 'SIM'|'NAO'} das linhas do quadro NR-6 que o formulário MARCOU."""
    out = {}
    corpo = _secao_form(form_text, 'EPIS FORNECIDOS') or (form_text or '')
    dentro = False
    for l in corpo.splitlines():
        if 'NR-6' in l.upper():
            dentro = True
            continue
        if dentro and re.match(r'^\s*#{1,4}\s', l):
            break
        if not dentro or not re.match(r'^\s*[-*•]', l):
            continue
        chave = _nr6_chave(l.split('—')[0])
        v = _checkbox_sim_nao(l)
        if chave and v:
            out[chave] = v
    return out


def gate_formulario(data, form_text):
    """[] se OK; senão lista de mensagens de erro. Gates 1.7 (neutralização) e
    1.8 (NR-6): o JSON não pode contrariar o que o perito MARCOU no formulário."""
    erros = []
    blocks = {(k[2:-2] if k.startswith('{{') and k.endswith('}}') else k): v
              for k, v in (data.get('blocks') or {}).items()}

    for marcador, v in neutralizacao_por_agente(form_text).items():
        if v == 'SIM' and marcador in blocks and caracteriza_insalubridade(blocks[marcador]):
            erros.append('  - %s: o formulário marca "Neutralizado pelo EPI: [X] Sim", mas o bloco '
                         'CARACTERIZA insalubridade. O checkbox do perito é a autoridade — corrija a '
                         'análise (ou o formulário, se a marcação estiver errada).' % marcador)

    nr6_form, nr6_json = nr6_do_form(form_text), (data.get('nr6') or {})
    for chave, val_form in nr6_form.items():
        val_json = str(nr6_json.get(chave, '') or '').strip().upper()
        if val_json and val_json != val_form:
            erros.append('  - nr6["%s"]: formulário diz %s, JSON diz %s. Linha marcada pelo perito não '
                         'é revista pelo redator — um NÃO estrutural derruba a neutralização do EPI e '
                         'muda o veredito.' % (chave, val_form, val_json))
    return erros


_INSAL = 'Descaracterizada a insalubridade.'
_PERIC = 'Descaracterizada a periculosidade.'
_NC = 'Não foi constatada, nas atividades exercidas pelo(a) Reclamante, exposição a %s nos termos do Anexo nº %s da NR-%s.'
# Item 3.1 — texto-padrão do Irineu quando NÃO há divergência fática (o modelo só manda
# {{DIVERGENCIAS_FATICAS}} em blocks quando HÁ divergência; ausente = esta frase).
DIVERGENCIA_FATICA_PADRAO = ('Durante a diligência pericial não houve divergência fática, '
                             'sendo as atividades do(a) Reclamante confirmadas pela Reclamada.')
# Honorários FIXOS do Irineu — cravados aqui (fonte única). O perito não digita mais valor/extenso;
# o build força estes valores no laudo, ignorando o que vier no JSON. Se o Irineu reajustar o padrão,
# muda-se SÓ estas 2 linhas (e o default do formulário em 01-extrator/montar_formulario.py).
HONORARIOS_VALOR_FIXO = '5.800,00'
HONORARIOS_EXTENSO_FIXO = 'Cinco mil e oitocentos reais'
ABSENT_ANALISE = {
    # NR-15
    'ANALISE_RUIDO_CONTINUO':   [_NC % ('ruído contínuo ou intermitente', '1', '15'), _INSAL],
    'ANALISE_RUIDO_IMPACTO':    [_NC % ('ruído de impacto', '2', '15'), _INSAL],
    'ANALISE_CALOR':            ['Não foi constatada, nas atividades exercidas pelo(a) Reclamante, exposição ao agente calor nos termos do Anexo nº 3 da NR-15.', _INSAL],
    'ANALISE_BAIXO_ILUMINAMENTO': ['O Anexo nº 4 da NR-15 (Iluminamento) foi revogado pela Portaria nº 3.751, de 23/11/1990, não sendo, portanto, objeto de análise no presente laudo.'],
    'ANALISE_RAD_IONIZANTES':   [_NC % ('radiações ionizantes', '5', '15'), _INSAL],
    'ANALISE_HIPERBARICAS':     ['Não foi constatado, nas atividades exercidas pelo(a) Reclamante, trabalho sob condições hiperbáricas nos termos do Anexo nº 6 da NR-15.', _INSAL],
    'ANALISE_RAD_NAO_IONIZANTES': ['Não foi constatada, nas atividades exercidas pelo(a) Reclamante, exposição a radiações não ionizantes em condições que se enquadrem no Anexo nº 7 da NR-15.', _INSAL],
    'ANALISE_VIBRACOES':        [_NC % ('vibrações', '8', '15'), _INSAL],
    'ANALISE_FRIO':             ['Não aplicável às atividades exercidas pelo(a) Reclamante, nos termos do Anexo nº 9 da NR-15.', _INSAL],
    'ANALISE_UMIDADE':          ['Não foi constatada, nas atividades exercidas pelo(a) Reclamante, exposição a condições de umidade que se enquadrem no Anexo nº 10 da NR-15.', _INSAL],
    'ANALISE_QUIM_QUANTITATIVOS': ['Não foi constatada, nas atividades exercidas pelo(a) Reclamante, exposição a agentes químicos sujeitos a avaliação quantitativa nos termos do Anexo nº 11 da NR-15.', _INSAL],
    'ANALISE_POEIRAS_MINERAIS': [_NC % ('poeiras minerais', '12', '15'), _INSAL],
    'ANALISE_QUIM_QUALITATIVOS': ['Não foi constatada, nas atividades exercidas pelo(a) Reclamante, exposição a agentes químicos sujeitos a avaliação qualitativa nos termos do Anexo nº 13 da NR-15.', _INSAL],
    'ANALISE_BENZENO':          ['Não foi constatada, nas atividades exercidas pelo(a) Reclamante, exposição a benzeno nos termos do Anexo nº 13-A da NR-15.', _INSAL],
    'ANALISE_AGENTES_BIOLOGICOS': [_NC % ('agentes biológicos', '14', '15'), _INSAL],
    # NR-16
    'ANALISE_PERIC_EXPLOSIVOS': [_NC % ('explosivos', '1', '16'), _PERIC],
    'ANALISE_PERIC_INFLAMAVEIS': [_NC % ('inflamáveis', '2', '16'), _PERIC],
    'ANALISE_PERIC_ROUBOS':     ['Não foi constatada, nas atividades exercidas pelo(a) Reclamante, exposição a situações de roubo ou outras espécies de violência física nos termos do Anexo nº 3 da NR-16.', _PERIC],
    'ANALISE_PERIC_ELETRICIDADE': ['Não foi constatado, nas atividades exercidas pelo(a) Reclamante, contato com instalações ou equipamentos elétricos energizados em condições que se enquadrem no Anexo nº 4 da NR-16.', _PERIC],
    'ANALISE_PERIC_MOTOCICLETA': ['Não foi constatado, nas atividades exercidas pelo(a) Reclamante, uso de motocicleta para o exercício do trabalho nos termos do Anexo nº 5 da NR-16.', _PERIC],
    'ANALISE_PERIC_RADIACOES':  ['Não foi constatada, nas atividades exercidas pelo(a) Reclamante, exposição a radiações ionizantes em condições de periculosidade nos termos da NR-16.', _PERIC],
}

# ---------------- helpers de parágrafo ----------------
def all_paragraphs(document):
    out = []
    def rec(tbls):
        for t in tbls:
            for row in t.rows:
                for cell in row.cells:
                    out.extend(cell.paragraphs); rec(cell.tables)
    out.extend(document.paragraphs); rec(document.tables)
    for sec in document.sections:
        out.extend(sec.header.paragraphs); out.extend(sec.footer.paragraphs)
    return out

def replace_scalar(document, mapping):
    for p in all_paragraphs(document):
        for run in p.runs:
            if '{{' in run.text:
                t = run.text
                for k, v in mapping.items():
                    if k in t:
                        t = t.replace(k, str(v))  # str(): scalar do JSON pode vir número (EPI_ANO etc) → sem str, crash
                run.text = t

def set_block(p, lines):
    rpr = None
    if p.runs:
        el = p.runs[0]._r.find(qn('w:rPr'))
        if el is not None:
            rpr = deepcopy(el)
            u = rpr.find(qn('w:u'))            # conteúdo de corpo nunca sublinhado
            if u is not None: rpr.remove(u)
    for r in list(p.runs):
        r._r.getparent().remove(r._r)
    def add_line(par, text):
        run = par.add_run(text)
        if rpr is not None:
            ex = run._r.find(qn('w:rPr'))
            if ex is not None: run._r.remove(ex)
            run._r.insert(0, deepcopy(rpr))
    if not lines: lines = ['']
    add_line(p, lines[0]); anchor = p
    for line in lines[1:]:
        new_p = deepcopy(p._p)
        for r in new_p.findall(qn('w:r')): new_p.remove(r)
        anchor._p.addnext(new_p)
        np = Paragraph(new_p, p._parent); add_line(np, line); anchor = np

def replace_blocks(document, blocks):
    for marker, lines in blocks.items():
        mk = '{{%s}}' % marker if not marker.startswith('{{') else marker
        for p in [p for p in all_paragraphs(document) if mk in p.text]:
            set_block(p, lines)

def fill_absent_analises(document):
    """Qualquer {{ANALISE_*}} que o modelo NÃO enviou (agente ausente) recebe a
    descaracterização-padrão por agente. Devolve a lista de marcadores preenchidos."""
    filled = []
    for p in all_paragraphs(document):
        m = re.search(r'\{\{(ANALISE_[A-Z0-9_]+)\}\}', p.text)
        if m and m.group(1) in ABSENT_ANALISE:
            set_block(p, ABSENT_ANALISE[m.group(1)])
            filled.append(m.group(1))
    return filled

def fill_default_divergencia(document):
    """Item 3.1: se o {{DIVERGENCIAS_FATICAS}} sobrou (o modelo não mandou em blocks,
    logo NÃO há divergência), preenche com o texto-padrão. Devolve True se preencheu."""
    for p in all_paragraphs(document):
        if '{{DIVERGENCIAS_FATICAS}}' in p.text:
            set_block(p, [DIVERGENCIA_FATICA_PADRAO])
            return True
    return False

def _norm_txt(s):
    """minúsculas sem acento (compara tipo de público)."""
    return ''.join(c for c in unicodedata.normalize('NFD', str(s))
                   if unicodedata.category(c) != 'Mn').lower().strip()

def _tc_text(tc):
    return ''.join(t.text or '' for t in tc.iter(qn('w:t')))

def _tc_set(tc, text):
    """seta o texto da 1ª <w:t> da célula (preserva a formatação do run); zera as demais."""
    nodes = list(tc.iter(qn('w:t')))
    if nodes:
        nodes[0].text = text
        for e in nodes[1:]:
            e.text = ''

def fill_banheiro(document, banheiro):
    """Item 3 — tabela de 'banheiros de grande circulação'. Só aparece quando o perito
    preencheu o bloco no formulário (JSON traz `banheiro` com `locais`): clona uma linha
    por local, soma os totais e preenche a rotatividade pelos tipos informados (remove os
    tipos em branco). SEM dados → REMOVE a tabela inteira (o gate de órfão exige que nenhum
    {{BANH_*}} sobre). Devolve True se preencheu, False se removeu/ausente."""
    tb = find_table(document, '{{BANH_LOCAL}}')
    if tb is None:
        return False
    tbl = tb._tbl
    WTR, WTC = qn('w:tr'), qn('w:tc')
    locais = (banheiro or {}).get('locais') or []
    if not locais:
        tbl.getparent().remove(tbl)     # caso não-banheiro: tabela some por completo
        return False
    # 1) uma linha por local (clona a linha-modelo {{BANH_LOCAL}}), somando os totais
    tmpl_tr = next(tr for tr in tbl.findall(WTR)
                   if any('{{BANH_LOCAL}}' in _tc_text(tc) for tc in tr.findall(WTC)))
    tot_b = tot_v = 0
    tot_ok = True
    for loc in locais:
        b, v, local = str(loc.get('banheiros', '')), str(loc.get('vasos', '')), str(loc.get('local', ''))
        new_tr = deepcopy(tmpl_tr)
        for tc in new_tr.findall(WTC):
            txt = _tc_text(tc)
            if '{{BANH_QTD}}' in txt:     _tc_set(tc, b)
            elif '{{BANH_VASOS}}' in txt: _tc_set(tc, v)
            elif '{{BANH_LOCAL}}' in txt: _tc_set(tc, local)
        tmpl_tr.addprevious(new_tr)
        try: tot_b += int(b); tot_v += int(v)
        except ValueError: tot_ok = False
    tmpl_tr.getparent().remove(tmpl_tr)
    # 2) totais (em branco se algum valor não for numérico)
    for tr in tbl.findall(WTR):
        for tc in tr.findall(WTC):
            txt = _tc_text(tc)
            if '{{BANH_QTD_TOTAL}}' in txt:   _tc_set(tc, str(tot_b) if tot_ok else '')
            elif '{{BANH_VASOS_TOTAL}}' in txt: _tc_set(tc, str(tot_v) if tot_ok else '')
    # 3) rotatividade: preenche os tipos informados; remove as linhas dos tipos em branco
    rot = {_norm_txt(k): str(v) for k, v in ((banheiro or {}).get('rotatividade') or {}).items() if str(v).strip()}
    marker_tipo = {'{{BANH_ROT_ALUNOS}}': 'alunos', '{{BANH_ROT_CLIENTES}}': 'clientes',
                   '{{BANH_ROT_FUNC}}': 'funcionarios', '{{BANH_ROT_PAC}}': 'pacientes'}
    kept = 0
    for tr in list(tbl.findall(WTR)):
        marker = next((m for m in marker_tipo
                       if any(m in _tc_text(tc) for tc in tr.findall(WTC))), None)
        if marker is None:
            continue
        tipo = marker_tipo[marker]
        if tipo in rot:
            for tc in tr.findall(WTC):
                if marker in _tc_text(tc): _tc_set(tc, rot[tipo])
            kept += 1
        else:
            tr.getparent().remove(tr)
    # nenhuma rotatividade informada -> remove o cabeçalho "Rotatividade" e a linha separadora
    if kept == 0:
        for tr in list(tbl.findall(WTR)):
            txts = [_tc_text(tc).strip() for tc in tr.findall(WTC)]
            if (txts and txts[0].startswith('Rotatividade')) or txts == ['', '', '']:
                tr.getparent().remove(tr)
    return True

def set_cell_text(cell, text, bold=None):
    p = cell.paragraphs[0]
    for extra in cell.paragraphs[1:]:
        extra._p.getparent().remove(extra._p)
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r._r.getparent().remove(r._r)
    else:
        p.add_run(text)
    if bold is not None:
        for r in p.runs: r.bold = bold

# ---------------- localizar tabelas por conteúdo ----------------
def find_table(doc, needle):
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                if needle in c.text:
                    return t
    return None


def _template_markers(doc):
    """Conjunto de marcadores {{NOME}} presentes no template (corpo+tabelas+header/footer)."""
    raw = '\n'.join(p.text for p in all_paragraphs(doc))
    return set(re.findall(r'\{\{([A-Z0-9_]+)\}\}', raw))


def _check_template_contract(path):
    """GUARD ANTI-DRIFT: valida que um template tem a estrutura que build_laudo.py exige.
    Pega tanto schema antigo (EPI de 3 linhas, imprescrito de 1 coluna) quanto marcadores
    de tipo errado. Devolve (problemas, tipo) — problemas vazio = template OK."""
    doc = docx.Document(path)
    m = _template_markers(doc)
    base = ntpath.basename(path)
    tipo = next((t for t, (b, _n15, _n16) in _TIPO_TEMPLATE.items() if b == base), None)
    probs = []
    # contrato comum a todos os tipos (schema atual)
    if not {'EPI_DESC', 'EPI_AGENTE'} <= m:
        probs.append('tabela de EPI fora da schema atual (falta {{EPI_DESC}}/{{EPI_AGENTE}})')
    if 'EPI_DESC_1' in m:
        probs.append('schema de EPI ANTIGO ({{EPI_DESC_1}}) — regerar do insal-peric')
    if not {'PERIODO_IMPRESCRITO_INICIO', 'PERIODO_IMPRESCRITO_TERMINO'} <= m:
        probs.append('falta imprescrito por função ({{PERIODO_IMPRESCRITO_INICIO/TERMINO}})')
    if 'PERIODO_IMPRESCRITO' in m:
        probs.append('imprescrito ANTIGO (coluna única {{PERIODO_IMPRESCRITO}})')
    it = find_table(doc, '{{FUNCAO}}'); ep = find_table(doc, '{{EPI_DESC}}'); nr6 = find_table(doc, 'NR-6 EQUIPAMENTO')
    if not it or len(it.rows[-1].cells) != 7:
        probs.append('tabela de identificação != 7 colunas (build espera 7)')
    if not ep or len(ep.rows[-1].cells) != 6:
        probs.append('tabela de EPI != 6 colunas (build espera 6)')
    if not nr6 or len(nr6.rows) != 9:
        probs.append('tabela NR-6 != 9 linhas (rowmap do build espera 9)')
    # marcadores de análise por tipo
    nr15 = [x for x in m if x.startswith('ANALISE_') and 'PERIC' not in x]
    peric = [x for x in m if x.startswith('ANALISE_PERIC_')]
    if tipo == 'insalubridade':
        if len(nr15) != 15: probs.append('esperava 15 marcadores NR-15, achou %d' % len(nr15))
        if peric: probs.append('template de INSALUBRIDADE não pode ter marcadores de periculosidade')
    elif tipo == 'periculosidade':
        if nr15: probs.append('template de PERICULOSIDADE não pode ter marcadores NR-15')
        if len(peric) != 6: probs.append('esperava 6 marcadores de periculosidade, achou %d' % len(peric))
    elif tipo == 'insal-peric':
        if len(nr15) != 15 or len(peric) != 6:
            probs.append('insal-peric deve ter 15 NR-15 + 6 periculosidade (achou %d+%d)' % (len(nr15), len(peric)))
    return probs, tipo

# ---------------- main ----------------
def _resolve_epi_paths(extras):
    """extras = caminhos soltos (qualquer ordem): .json=dicionário, .sqlite=CAEPI,
    diretório=base (deriva 04-EPIs/CA-dicionario.json e 04-EPIs/caepi.sqlite)."""
    caepi_p = dicio_p = None
    for a in extras:
        if not a:
            continue
        if os.path.isdir(a):
            caepi_p = caepi_p or os.path.join(a, '04-EPIs', 'caepi.sqlite')
            dicio_p = dicio_p or os.path.join(a, '04-EPIs', 'CA-dicionario.json')
        elif a.endswith(('.sqlite', '.db')):
            caepi_p = a
        elif a.endswith('.json'):
            dicio_p = a
    # fallback BUNDLED — no Cowork o bash não vê o Drive; se o resolvido não existe, usa o do plugin
    if not (caepi_p and os.path.isfile(caepi_p)):
        b = os.path.join(_BUNDLED_EPI, 'caepi.sqlite')
        caepi_p = b if os.path.isfile(b) else caepi_p
    if not (dicio_p and os.path.isfile(dicio_p)):
        b = os.path.join(_BUNDLED_EPI, 'CA-dicionario.json')
        dicio_p = b if os.path.isfile(b) else dicio_p
    return caepi_p, dicio_p


def build(template_path, data_path, out_path, *epi_paths, form_path=None):
    data = json.load(open(data_path, encoding='utf-8'))
    caepi_p, dicio_p = _resolve_epi_paths(list(epi_paths) + [data.get('ca_dicionario_path'), data.get('caepi_path')])
    cadict = _load_ca_dict(dicio_p)
    caepi = _open_caepi(caepi_p)
    doc = docx.Document(_resolve_template(template_path))
    warnings = []

    # GATE DE TIPO — ANTES de qualquer mutação: se o template não casa com o tipo do
    # laudo, falha dura SEM gravar arquivo (nada de laudo defeituoso na mão).
    tipo_errs, _tipo, _intent = _gate_tipo(doc, template_path, data.get('tipo_laudo'))
    if tipo_errs:
        print('\n❌ LAUDO NÃO GERADO — template incompatível com o tipo do laudo:')
        for e in tipo_errs:
            print('  - ' + e)
        print('   Nenhum arquivo foi salvo. Corrija (JSON/comando) e rode de novo.')
        return False
    if _tipo is None:
        warnings.append('tipo_laudo AUSENTE no JSON — declare "insalubridade"/"periculosidade"/'
                        '"insal-peric" (do ▶ TIPO DE LAUDO do formulário) p/ o gate cruzar com o formulário')

    # GATE DE CONSISTÊNCIA JSON × FORMULÁRIO (aborta ANTES de gravar — os dois piores
    # erros que o gate de órfão NÃO pega, porque fill_absent_analises os mascara):
    #  (a) MARCADOR ANALISE_* INVÁLIDO — chave fora do padrão (ANALISE_RUIDO em vez de
    #      ANALISE_RUIDO_CONTINUO, ou ANALISE_QUIMICOS_ACIDOS inexistente): descartada em
    #      silêncio por replace_blocks e o marcador REAL é preenchido pela descaracterização.
    #  (b) AGENTE PRESENTE OMITIDO — agente [Presente] no formulário (agentes_presentes)
    #      que o modelo esqueceu de escrever em blocks: fill_absent_analises o
    #      descaracteriza. Nos dois casos: CARACTERIZADO vira descaracterizado, sem órfão.
    _valid_analise = set(ABSENT_ANALISE)
    _mk = lambda k: k[2:-2] if k.startswith('{{') and k.endswith('}}') else k
    _emitted = {_mk(k) for k in (data.get('blocks') or {})}
    bad_markers = sorted({k for k in _emitted if k.startswith('ANALISE_') and k not in _valid_analise})
    _present = [_mk(k) for k in (data.get('agentes_presentes') or [])]
    present_bad = sorted({a for a in _present if a not in _valid_analise})
    present_faltando = sorted({a for a in _present if a in _valid_analise and a not in _emitted})
    if bad_markers or present_bad or present_faltando:
        print('\n❌ LAUDO NÃO GERADO — inconsistência entre o formulário e as análises do JSON:')
        for k in bad_markers:
            print('  - chave inválida em blocks: {{%s}} — não existe, seria DESCARTADA em silêncio '
                  '(agente caracterizado viraria descaracterizado)' % k)
        for a in present_faltando:
            print('  - agente PRESENTE no formulário sem análise no JSON: %s — seria DESCARACTERIZADO em silêncio' % a)
        for a in present_bad:
            print('  - nome inválido em agentes_presentes: %s (não é um ANALISE_* canônico)' % a)
        print('   ANALISE_* válidos: %s' % ', '.join(sorted(_valid_analise)))
        print('   Nenhum arquivo foi salvo. Corrija o JSON e rode de novo.')
        return False

    # GATES 1.7/1.8 — o formulário é a autoridade sobre o FATO (só com --form)
    if form_path:
        try:
            _form_text = open(form_path, encoding='utf-8').read()
        except OSError as e:
            print('\n❌ LAUDO NÃO GERADO — não li o --form: %s' % e)
            return False
        _erros = gate_formulario(data, _form_text)
        if _erros:
            print('\n❌ LAUDO NÃO GERADO — o JSON contraria o formulário do perito:')
            for e in _erros:
                print(e)
            print('   O checkbox do perito vence texto-padrão, paradigma e caso anterior:')
            print('   eles dão a moldura, nunca o fato. Divergência é dúvida — pergunte ao perito.')
            print('   Nenhum arquivo foi salvo.')
            return False
    else:
        warnings.append('sem --form: gates de neutralização e NR-6 NÃO verificados')

    # blocos (multi-parágrafo) primeiro, depois escalares
    replace_blocks(doc, data.get('blocks', {}))
    # agentes AUSENTES que o modelo não enviou -> descaracterização-padrão automática
    auto = fill_absent_analises(doc)
    if auto:
        print('Agentes ausentes preenchidos pelo script (%d): %s' % (len(auto), ', '.join(a.replace('ANALISE_', '') for a in auto)))
    # item 3.1: sem divergência informada -> texto-padrão "não houve divergência fática..."
    if fill_default_divergencia(doc):
        print('Item 3.1 (Divergências Fáticas): sem divergência no JSON — usado o texto-padrão.')

    scalars = dict(data.get('scalars', {}))
    # honorários FIXOS — sobrepõe qualquer valor do JSON/formulário (perito não digita mais)
    scalars['HONORARIOS_VALOR'] = HONORARIOS_VALOR_FIXO
    scalars['HONORARIOS_EXTENSO'] = HONORARIOS_EXTENSO_FIXO
    # anos do EPI no cabeçalho
    anos = data.get('epi', {}).get('anos', [])
    for i in range(3):
        scalars['{{EPI_ANO_%d}}' % (i+1)] = anos[i] if i < len(anos) else ''
    # legendas de foto são sempre manuais -> limpar marcador
    for i in range(1, 11):
        scalars.setdefault('{{LEGENDA_FOTO_%d}}' % i, '')
    # normalizar chaves para {{...}}
    scalars = {(k if k.startswith('{{') else '{{%s}}' % k): v for k, v in scalars.items()}
    replace_scalar(doc, scalars)

    # tabela de identificação
    t0 = find_table(doc, '{{FUNCAO}}')
    if t0 is not None and data.get('identificacao'):
        tmpl = t0.rows[-1]._tr
        for row_vals in data['identificacao']:
            new_tr = deepcopy(tmpl); t0._tbl.append(new_tr)
            row = t0.rows[-1]
            for ci, val in enumerate(row_vals[:len(row.cells)]):
                set_cell_text(row.cells[ci], str(val), bold=False)
        t0._tbl.remove(tmpl)

    # guard determinístico de EPI ANTES de montar a tabela: C.A.=chave (lookup no
    # dicionário), creme->An.13 (regra absoluta), flaga o resto. Muta as linhas.
    epi_fixes, epi_flags, epi_naocat = epi_guard(data.get('epi', {}).get('linhas') or [], cadict, caepi)

    # tabela de EPI (resumo por agente) — tolerante: cada linha preenche TODAS as células
    # (desc, agente, ca, v1, v2, v3); campo faltante -> '' (nunca deixa {{EPI_*}} residual).
    t2 = find_table(doc, '{{EPI_DESC}}')
    if t2 is not None:
        tmpl = t2.rows[-1]._tr
        for row_vals in (data.get('epi', {}).get('linhas') or []):
            new_tr = deepcopy(tmpl); t2._tbl.append(new_tr)
            row = t2.rows[-1]
            for ci in range(len(row.cells)):
                val = row_vals[ci] if ci < len(row_vals) else ''
                set_cell_text(row.cells[ci], str(val))
        t2._tbl.remove(tmpl)   # remove a linha-modelo (some os {{EPI_*}} mesmo sem dados)

    # NR-6
    t3 = find_table(doc, 'NR-6 EQUIPAMENTO')
    if t3 is not None and data.get('nr6'):
        rowmap = {'ficha':3,'ca':4,'treinamento':5,'adequado':6,'frequencia':7,'fiscalizacao':8}
        for key, ri in rowmap.items():
            v = (data['nr6'].get(key) or '').strip().upper()
            col = 1 if v == 'SIM' else (2 if v in ('NAO','NÃO') else None)
            if col is not None and ri < len(t3.rows):
                set_cell_text(t3.rows[ri].cells[col], 'X')

    # tabela de vibração (substitui @@TABELA_VIBRACAO@@)
    if data.get('vibracao'):
        build_vibracao_table(doc, data['vibracao'])

    # item 3 — tabela de banheiros de grande circulação (só se o perito preencheu o bloco;
    # sem dados a tabela é REMOVIDA por completo, senão sobraria {{BANH_*}} órfão)
    if fill_banheiro(doc, data.get('banheiro')):
        print('Item 3: tabela de banheiros de grande circulação preenchida (%d local/ais).'
              % len(data['banheiro'].get('locais', [])))

    # ---- validações ----
    full = '\n'.join(p.text for p in all_paragraphs(doc))
    # GATE DE MARCADOR ÓRFÃO (paridade com o squad laudo-pericial v2.1): qualquer
    # {{CAMPO}} que sobrou = laudo incompleto. Mantido FORA de `warnings` para ganhar
    # bloco próprio e explícito no relatório (não enterrado entre avisos de EPI) e
    # ainda assim bloquear a saída (exit 2) — só avisar não basta, senão o laudo
    # incompleto é declarado concluído. all_paragraphs cobre corpo + tabelas
    # (aninhadas) + header/footer (mais amplo que o scan_orphans do squad).
    orphans = sorted(set(re.findall(r'\{\{[^}]+\}\}', full)))
    perito = data.get('perito_nome', 'Irineu de Freitas Branco Junior')
    if perito not in full:
        warnings.append('IDENTIDADE: nome do perito (%s) não encontrado no documento' % perito)
    for bad in list(data.get('nomes_proibidos', [])) + ['@@TABELA']:
        if bad in full:
            warnings.append('VAZAMENTO: "%s" presente no documento' % bad)

    doc.save(out_path)

    # --- relatório autossuficiente: tudo que dispensa reabrir o .docx ---
    n_paras = len([p for p in all_paragraphs(doc) if p.text.strip()])
    n_tbls = len(doc.tables)
    ident_rows = len(data.get('identificacao', []) or [])
    epi_rows = len(data.get('epi', {}).get('linhas', []) or [])
    vib_rows = len(data.get('vibracao', []) or [])
    print('OK ->', out_path)
    print('CONTEÚDO: %d parágrafos não-vazios | %d tabelas (identificação=%d linhas, EPI=%d linhas, vibração=%d linhas)'
          % (n_paras, n_tbls, ident_rows, epi_rows, vib_rows))
    if auto:
        print('Agentes AUSENTES preenchidos pelo script: %d/%d' % (len(auto), 21))
    if orphans:
        print('\n❌ LAUDO INCOMPLETO — %d marcador(es) órfão(s): %s' % (len(orphans), ', '.join(orphans)))
        print('   Corrigir o JSON (campos faltando) e re-gerar — NÃO entregar este .docx.')
    if warnings:
        print('\n⚠ AVISOS (corrija o JSON e rode de novo):')
        for w in warnings: print('  -', w)
    if epi_fixes:
        print('\n🔧 EPI — CLASSIFICADO POR C.A./REGRA (C.A.=chave; nome comercial não classifica):')
        for trecho, msg in epi_fixes:
            print('  - %s → %s' % (trecho, msg))
    if epi_flags:
        print('\n🚩 EPI — CONFERIR CLASSIFICAÇÃO PELO C.A.:')
        for trecho, msg in epi_flags:
            print('  - %s → %s' % (trecho, msg))
    if epi_naocat:
        print('\n📇 EPI — C.A. NÃO CATALOGADOS (adicione ao CA-dicionario.json): %s' % ', '.join(epi_naocat))
    if not orphans and not warnings and not bad_markers and not epi_flags:
        if epi_fixes:
            print('\n✅ DOCUMENTO GERADO — creme(s) auto-corrigido(s) para An.13 (acima); nada mais pendente.')
        else:
            print('✅ VALIDAÇÃO OK: sem marcador residual, identidade do perito presente, sem vazamento.')
        print('✅ verificação encerrada. NÃO reabra/dumpe o .docx: ele é render determinístico do JSON já conferido.')
    elif not orphans and not warnings and not bad_markers and epi_flags:
        print('\n⚠ DOCUMENTO GERADO, mas há classificação(ões) de EPI a confirmar pelo C.A. acima — revise ANTES de assinar.')
    return not warnings and not orphans and not bad_markers

def build_vibracao_table(doc, linhas):
    ph = None
    for p in doc.paragraphs:
        if '@@TABELA_VIBRACAO@@' in p.text:
            ph = p; break
    if ph is None:
        return
    header = ('Equipamento', 'AREN\nLT = 1,1 m/s²', 'VDVR\nLT = 21,0')
    rows = [header] + [tuple(r) for r in linhas]
    tbl = doc.add_table(rows=len(rows), cols=3)
    tbl.style = doc.styles['Table Grid']
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ri, rowvals in enumerate(rows):
        for ci, val in enumerate(rowvals):
            cell = tbl.rows[ri].cells[ci]
            para = cell.paragraphs[0]; para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for li, line in enumerate(str(val).split('\n')):
                run = para.add_run(line); run.font.name = 'Arial'; run.font.size = Pt(10.5)
                if ri == 0:
                    run.bold = True; run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                if li < len(str(val).split('\n')) - 1:
                    run.add_break()
            if ri == 0:
                tcPr = cell._tc.find(qn('w:tcPr'))
                if tcPr is None:
                    tcPr = OxmlElement('w:tcPr'); cell._tc.insert(0, tcPr)
                shd = OxmlElement('w:shd'); shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), '4472C4')
                tcPr.append(shd)
    ph._p.addprevious(tbl._tbl)
    ph._p.getparent().remove(ph._p)

# ---------------- guard de classificação de EPI por agente ----------------
# Mesma regra dura do extrator (check_epi.py): a falha perigosa é classificar o
# EPI pelo NOME COMERCIAL (creme "Luz Negra" -> UV/An.7), tirando-o do An.13.
_EPI_RAD = ('radiaç', 'radiac', 'rni', 'não ioniz', 'nao ioniz', 'ultraviolet',
            'luz negra', 'an.7', 'an. 7', 'anexo 7')
_EPI_QUIM = ('quím', 'quim', 'an.13', 'an. 13', 'anexo 13', 'dérm', 'derm',
             'óleo', 'oleo', 'graxa', 'álcali', 'alcali', 'solvente')
_EPI_UMID = ('umidade', 'an.10', 'an. 10', 'anexo 10')
_EPI_MASK = ('máscara', 'mascara', 'lente', 'viseira', 'escudo', 'facial',
             'solda', 'soldad', 'capuz')

def _load_ca_dict(path):
    if not path:
        return {}
    try:
        with open(path, encoding='utf-8') as f:
            raw = json.load(f)
    except Exception:
        return {}
    out = {}
    for k, v in raw.items():
        if str(k).startswith('_'):
            continue
        key = re.sub(r'\D', '', str(k))
        if key:
            out[key] = v
    return out


def _caepi_lookup(con, ca):
    """Devolve (known, agente): known=C.A. existe na base (mesmo sem agente NR-15)."""
    if not con or not ca:
        return False, None
    try:
        r = con.execute('SELECT agente FROM ca WHERE ca=?', (ca,)).fetchone()
    except Exception:
        return False, None
    if r is None:
        return False, None
    return True, (r[0] or None)


def _open_caepi(path):
    if path and os.path.exists(path):
        try:
            # as_uri() (file:///C:/... percent-encoded) — 'file:%s' com path Windows
            # (C:\...) sai fora do formato de URI que o SQLite documenta
            return sqlite3.connect(Path(path).resolve().as_uri() + '?mode=ro', uri=True)
        except Exception:
            return None
    return None


def epi_guard(linhas, cadict=None, caepi=None):
    """linhas = [[desc, agente, ca, ...], ...]. C.A. é a CHAVE: override curado
    (cadict) vence; depois base oficial CAEPI; sem nenhum → regra absoluta
    creme/pomada=An.13 (exceto protetor solar). MUTA as linhas.
    Devolve (fixes, flags, nao_catalogados)."""
    cadict = cadict or {}
    fixes, flags, nao_cat = [], [], []
    for row in linhas:
        if len(row) < 2:
            continue
        desc = str(row[0]).lower()
        ag = str(row[1]).lower()
        trecho = ('%s · %s' % (str(row[0]), str(row[1]))).strip(' ·')[:120]
        ca = re.sub(r'\D', '', str(row[2])) if len(row) > 2 else ''

        # (1) LOOKUP por C.A. — override curado vence; depois CAEPI oficial
        agente = src = None
        known = False
        if ca and ca in cadict:
            agente, src, known = cadict[ca].get('agente'), 'dicionário', True
        elif ca:
            known, hit = _caepi_lookup(caepi, ca)
            if hit:
                agente, src = hit, 'CAEPI'
        if agente:
            if str(row[1]) != agente:
                fixes.append((trecho, 'C.A. %s → %s [%s]' % (ca, agente, src)))
                row[1] = agente
            continue
        # C.A. conhecido na base sem agente NR-15 (botina, óculos, luva mecânica…) =
        # não-neutralizador → silêncio (NÃO é não catalogado, NÃO aplica heurística).
        if known:
            continue
        if ca:
            nao_cat.append(ca)

        # (2) C.A. desconhecido: regra absoluta creme/pomada=An.13 (exceto solar) + flags
        has_rad = any(t in ag for t in _EPI_RAD)
        has_quim = any(t in ag for t in _EPI_QUIM)
        is_creme = 'creme' in desc or 'pomada' in desc
        is_solar = 'solar' in desc or 'solar' in ag
        if is_creme and not is_solar and (has_rad or not has_quim):
            row[1] = 'Químico dérmico (An.13)'
            fixes.append((trecho, 'creme/pomada → Químico dérmico (An.13) [regra absoluta]'))
            ag = row[1].lower(); has_rad = False; has_quim = True
        has_umid = any(t in ag for t in _EPI_UMID)
        is_capa = 'capa' in desc or 'impermeáv' in desc or 'impermeav' in desc
        is_mask = any(t in desc for t in _EPI_MASK)
        if has_rad and not is_mask and not is_creme and not is_solar:
            flags.append((trecho, 'EPI em radiação/UV sem ser máscara/lente/escudo de solda — nome comercial ≠ agente. Confira o C.A.'))
        if is_capa and has_quim and not has_umid:
            flags.append((trecho, 'capa/impermeável como químico — vestimenta impermeável protege UMIDADE (An.10). Confirme.'))
    def dedup(xs):
        seen, out = set(), []
        for x in xs:
            if x not in seen:
                seen.add(x); out.append(x)
        return out
    return dedup(fixes), dedup(flags), dedup(nao_cat)

if __name__ == '__main__':
    # escape-hatch de auto-conferência: lista os ANALISE_* válidos sem "ler o código"
    if len(sys.argv) >= 2 and sys.argv[1] in ('--list-markers', '--markers'):
        print('Marcadores ANALISE_* válidos (%d):' % len(ABSENT_ANALISE))
        for m in ABSENT_ANALISE: print('  ' + m)
        sys.exit(0)
    # verificação avulsa: quais templates estão bundled (fallback do Cowork)
    if len(sys.argv) >= 2 and sys.argv[1] == '--list-templates':
        found = sorted(_list_bundled_templates())
        esperados = ['template-insalubridade.docx', 'template-periculosidade.docx', 'template-insal-peric.docx']
        print('Templates BUNDLED em %s:' % _BUNDLED_TEMPLATES)
        for e in esperados:
            print('  [%s] %s' % ('OK' if e in found else 'FALTA', e))
        extra = [f for f in found if f not in esperados]
        if extra:
            print('  (outros: %s)' % ', '.join(extra))
        sys.exit(0)
    # GUARD ANTI-DRIFT: valida a ESTRUTURA dos 3 templates bundled contra o contrato
    # do build. Rodar depois de qualquer edição de template — pega deriva antes do laudo.
    if len(sys.argv) >= 2 and sys.argv[1] == '--check-templates':
        esperados = ['template-insalubridade.docx', 'template-periculosidade.docx', 'template-insal-peric.docx']
        all_ok = True
        print('Contrato estrutural dos templates BUNDLED:')
        for e in esperados:
            p = os.path.join(_BUNDLED_TEMPLATES, e)
            if not os.path.isfile(p):
                print('  [FALTA] %s' % e); all_ok = False; continue
            probs, tipo = _check_template_contract(p)
            if probs:
                all_ok = False
                print('  [FALHA] %s' % e)
                for x in probs: print('        - ' + x)
            else:
                print('  [OK]    %s (%s)' % (e, tipo))
        print('\n%s' % ('✅ os 3 templates cumprem o contrato do build.' if all_ok
                        else '❌ há template fora do contrato — NÃO rodar laudos até corrigir.'))
        sys.exit(0 if all_ok else 2)
    # --form <path> em qualquer posição (retrocompatível com a CLI posicional)
    _argv, _form = list(sys.argv[1:]), None
    if '--form' in _argv:
        _i = _argv.index('--form')
        if _i + 1 >= len(_argv):
            print('erro: --form exige o caminho do formulário'); sys.exit(1)
        _form = _argv[_i + 1]
        del _argv[_i:_i + 2]
    if len(_argv) < 3:
        print('uso: python3 build_laudo.py <template.docx> <laudo-data.json> <saida.docx> [<caepi.sqlite>] [<CA-dicionario.json>] [<base_dir>] [--form <formulario.md>]')
        print('     --form ativa os gates 1.7/1.8: o JSON não pode contrariar checkbox marcado pelo perito')
        print('     python3 build_laudo.py --list-markers | --list-templates | --check-templates'); sys.exit(1)
    ok = build(_argv[0], _argv[1], _argv[2], *_argv[3:], form_path=_form)
    sys.exit(0 if ok else 2)
