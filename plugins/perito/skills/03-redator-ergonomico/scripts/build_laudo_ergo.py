# -*- coding: utf-8 -*-
"""
build_laudo_ergo.py — monta o laudo ERGONÔMICO (.docx) a partir de:
  - um JSON de conteúdo (dados do processo: partes, vistoria, atividades, quesitos, cargos)
  - a PLANILHA preenchida do caso (.xlsx): formulações, níveis, qualificação e as 5 tabelas de checklist

Uso: python3 build_laudo_ergo.py <template-ergonomico.docx> <laudo-data.json> <planilha.xlsx> <saida.docx>

Divisão: o MODELO produz só o JSON (dados do processo). O SCRIPT lê a planilha
(escores/formulações/tabelas — determinístico, nunca recalcula) e monta o .docx.
"""
import sys, json, re, os
from copy import deepcopy

# --- auto-provisionamento de dependências (sandbox efêmero do Cowork) ---
def _ensure(pkgs):
    import importlib.util, subprocess, sys as _sys
    falta = [pip for mod, pip in pkgs if importlib.util.find_spec(mod) is None]
    if falta:
        cmd = [_sys.executable, '-m', 'pip', 'install', *falta]
        if _sys.platform.startswith('linux'):
            cmd.append('--break-system-packages')
        subprocess.run(cmd, check=False)
_ensure([('docx', 'python-docx'), ('openpyxl', 'openpyxl')])

import docx, openpyxl
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from docx.shared import Pt, RGBColor, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ---------- helpers de docx ----------
def all_paragraphs(document):
    out = []
    def rec(tbls):
        for t in tbls:
            for row in t.rows:
                for cell in row.cells:
                    out.extend(cell.paragraphs); rec(cell.tables)
    out.extend(document.paragraphs); rec(document.tables)
    for sec in document.sections:
        out.extend(sec.header.paragraphs); out.extend(sec.footer.paragraphs)
    return out

def replace_scalar(document, mapping):
    for p in all_paragraphs(document):
        for run in p.runs:
            if '{{' in run.text:
                t = run.text
                for k, v in mapping.items():
                    if k in t: t = t.replace(k, str(v))
                run.text = t

def set_block(p, lines, justify=False):
    rpr = None
    if p.runs:
        el = p.runs[0]._r.find(qn('w:rPr'))
        if el is not None:
            rpr = deepcopy(el)
            u = rpr.find(qn('w:u'))
            if u is not None: rpr.remove(u)
    for r in list(p.runs): r._r.getparent().remove(r._r)
    if not lines: lines = ['']
    def add_line(par, text):
        run = par.add_run(text)
        if rpr is not None:
            ex = run._r.find(qn('w:rPr'))
            if ex is not None: run._r.remove(ex)
            run._r.insert(0, deepcopy(rpr))
    add_line(p, lines[0])
    if justify: p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    anchor = p
    for line in lines[1:]:
        new_p = deepcopy(p._p)
        for r in new_p.findall(qn('w:r')): new_p.remove(r)
        anchor._p.addnext(new_p)
        np = Paragraph(new_p, p._parent); add_line(np, line)
        if justify: np.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        anchor = np

def replace_blocks(document, blocks, justify=False):
    for marker, lines in blocks.items():
        mk = marker if marker.startswith('{{') else '{{%s}}' % marker
        for p in [p for p in all_paragraphs(document) if mk in p.text]:
            set_block(p, lines, justify=justify)

def set_cell_text(cell, text, bold=None):
    p = cell.paragraphs[0]
    for extra in cell.paragraphs[1:]: extra._p.getparent().remove(extra._p)
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]: r._r.getparent().remove(r._r)
    else:
        p.add_run(text)
    for r in p.runs:
        r.font.name = 'Arial'; r.font.size = Pt(10.5)
        if bold is not None: r.bold = bold

def find_table(doc, needle):
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                if needle in c.text: return t
    return None

def insert_table_at_anchor(doc, anchor_text, header, rows, widths=None, justify_cols=()):
    """header: lista (será MAIÚSCULA). widths: larguras em polegadas por coluna.
       justify_cols: índices de coluna cujo corpo é justificado (texto longo)."""
    ph = None
    for p in doc.paragraphs:
        if anchor_text in p.text: ph = p; break
    if ph is None: return False
    header = [str(h).upper() for h in header]      # cabeçalho em MAIÚSCULAS
    data = [header] + rows
    ncol = len(header)
    tbl = doc.add_table(rows=len(data), cols=ncol)
    tbl.style = doc.styles['Table Grid']
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    for ri, rowvals in enumerate(data):
        for ci in range(ncol):
            val = rowvals[ci] if ci < len(rowvals) else ''
            cell = tbl.rows[ri].cells[ci]
            if widths and ci < len(widths):
                cell.width = Inches(widths[ci])
            para = cell.paragraphs[0]
            # coluna larga (justify_cols) = justificada; demais (e cabeçalho) = centralizadas
            para.alignment = (WD_ALIGN_PARAGRAPH.JUSTIFY if (ri > 0 and ci in justify_cols)
                              else WD_ALIGN_PARAGRAPH.CENTER)
            run = para.add_run(str(val)); run.font.name = 'Arial'; run.font.size = Pt(10.5)
            if ri == 0:
                run.bold = True; run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                tcPr = cell._tc.find(qn('w:tcPr'))
                if tcPr is None:
                    tcPr = OxmlElement('w:tcPr'); cell._tc.insert(0, tcPr)
                shd = OxmlElement('w:shd'); shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), '4472C4')
                tcPr.append(shd)
    ph._p.addprevious(tbl._tbl)
    ph._p.getparent().remove(ph._p)
    return True

# ---------- leitura da planilha ----------
def _find_label_row(ws, label):
    for r in range(1, ws.max_row + 1):
        for c in range(1, 10):
            v = ws.cell(row=r, column=c).value
            if isinstance(v, str) and label in v:
                return r, c
    return None, None

def read_eval_aba(ws, pcol):
    """Retorna (linhas[[item,desc,pontos]], total, interpretacao)."""
    rows = []
    for r in range(7, ws.max_row + 1):
        b = ws.cell(row=r, column=2).value
        c = ws.cell(row=r, column=3).value
        p = ws.cell(row=r, column=pcol).value
        if isinstance(b, str) and b.strip().upper().startswith('RESUMO'):
            break
        if b is None and c is None and p is None:
            continue
        rows.append([('' if b is None else str(b)).strip(),
                     ('' if c is None else str(c)).strip(),
                     ('' if p is None else str(p)).strip()])
    # total: linha do rótulo "Total de Pontos:", valor na coluna de pontos
    tr, tc = _find_label_row(ws, 'Total de Pontos')
    total = ws.cell(row=tr, column=pcol).value if tr else None
    # interpretação: linha do rótulo "Interpretação:", valor em col D (4) ou adjacente
    ir, ic = _find_label_row(ws, 'Interpretação')
    interp = ws.cell(row=ir, column=4).value if ir else None
    if interp is None and ir:
        interp = ws.cell(row=ir, column=ic + 1).value
    return rows, ('' if total is None else str(total)), ('' if interp is None else str(interp))

def read_cond_aba(ws):
    """Abas 2 e 7: [[condição, X-SIM, X-NÃO]]."""
    rows = []
    for r in range(7, ws.max_row + 1):
        b = ws.cell(row=r, column=2).value
        if b is None or (isinstance(b, str) and b.strip().upper().startswith('RESUMO')):
            if b is None: continue
            break
        sim = ws.cell(row=r, column=3).value
        nao = ws.cell(row=r, column=4).value
        rows.append([str(b).strip(),
                     'X' if (sim and str(sim).strip()) else '',
                     'X' if (nao and str(nao).strip()) else ''])
    return rows

def derive_level(text, terms):
    up = (text or '').upper()
    for term in terms:               # do mais longo p/ o mais curto
        if term in up: return term
    return ''

def read_planilha(path):
    wb = openpyxl.load_workbook(path, data_only=True)
    L = wb['9-LAUDO']
    f = {k: (L[cell].value or '') for k, cell in
         {'BM': 'C4', 'MS': 'C6', 'CV': 'C8', 'C10': 'C10', 'C12': 'C12'}.items()}
    bm = derive_level(f['BM'], ['EXCELENTE', 'PÉSSIMA', 'RUIM', 'RAZOÁVEL', 'BOA'])
    ms = derive_level(f['MS'], ['ALTÍSSIMA', 'ALTA', 'MÉDIA', 'BAIXA'])
    cv = derive_level(f['CV'], ['ALTÍSSIMA', 'ALTA', 'MÉDIA', 'BAIXA'])
    def marcou(txt):
        t = (txt or '').strip()
        return bool(t) and 'NÃO HOUVE' not in t.upper()
    inadeq = (bm in ('RUIM', 'PÉSSIMA') or ms in ('ALTA', 'ALTÍSSIMA') or
              cv in ('ALTA', 'ALTÍSSIMA') or marcou(f['C10']) or marcou(f['C12']))
    qualif = 'ERGONOMICAMENTE INADEQUADAS' if inadeq else 'ERGONOMICAMENTE ADEQUADAS'
    tab = {}
    bmr, bmt, bmi = read_eval_aba(wb['3-AV.BM'], 7)
    msr, mst, msi = read_eval_aba(wb['5-AV.MS'], 7)
    cvr, cvt, cvi = read_eval_aba(wb['8-AV.CV'], 6)
    tab['BM'] = (bmr, bmt, bmi); tab['MS'] = (msr, mst, msi); tab['CV'] = (cvr, cvt, cvi)
    tab['COND'] = read_cond_aba(wb['2-COND.ERG.RUIM'])
    tab['ASECV'] = read_cond_aba(wb['7-ASECV'])
    return dict(formulacoes=f, niveis=dict(BM=bm, MS=ms, CV=cv),
                qualificacao=qualif, tabelas=tab)

# ---------- montagem ----------
def _resolve_template(template_path):
    """Fallback BUNDLED — no Cowork o bash NÃO enxerga o Drive; cai no template do plugin
    (assets/templates/). Nativo (bash vê o Drive) usa o caminho vivo."""
    if template_path and os.path.isfile(template_path):
        return template_path
    here = os.path.dirname(os.path.abspath(__file__))
    tdir = os.path.join(here, '..', 'assets', 'templates')
    cand = os.path.join(tdir, os.path.basename(template_path or ''))
    if os.path.isfile(cand):
        print('ℹ️  template do Drive inacessível (bash do Cowork) — usando o BUNDLED: %s'
              % os.path.basename(cand))
        return cand
    raise SystemExit('template não encontrado: nem "%s" nem bundled em %s' % (template_path, tdir))


def build(template, data_path, planilha, out_path):
    data = json.load(open(data_path, encoding='utf-8'))
    pl = read_planilha(planilha)
    doc = docx.Document(_resolve_template(template))
    warnings = []

    for dim, lvl in pl['niveis'].items():
        if not lvl: warnings.append('NÍVEL %s não derivado da formulação (célula vazia?)' % dim)

    # blocos do JSON (dados do processo) — alinhamento do template
    replace_blocks(doc, dict(data.get('blocks', {})))
    # formulações da planilha (item 8) — JUSTIFICADAS
    replace_blocks(doc, {
        '{{FORMULACAO_BIOMECANICA}}': [pl['formulacoes']['BM']],
        '{{FORMULACAO_MEMBROS_SUPERIORES}}': [pl['formulacoes']['MS']],
        '{{FORMULACAO_COLUNA_VERTEBRAL}}': [pl['formulacoes']['CV']],
        '{{FORMULACAO_COND_ERGONOMICA_EXTREMA}}': [pl['formulacoes']['C10']],
        '{{FORMULACAO_SITUACOES_EXTREMAS_COLUNA}}': [pl['formulacoes']['C12']],
    }, justify=True)

    # escalares: dados do processo (JSON) + níveis/qualificação (planilha)
    scal = dict(data.get('scalars', {}))
    scal['COND_BIOMECANICA'] = pl['niveis']['BM']
    scal['EXIG_MEMBROS_SUPERIORES'] = pl['niveis']['MS']
    scal['EXIG_COLUNA_VERTEBRAL'] = pl['niveis']['CV']
    scal['QUALIFICACAO_FINAL'] = pl['qualificacao']
    for i in range(1, 9):
        scal.setdefault('LEGENDA_FOTO_%d' % i, '')
    scal = {(k if k.startswith('{{') else '{{%s}}' % k): v for k, v in scal.items()}
    replace_scalar(doc, scal)

    # tabela de cargos (Item 2)
    t0 = find_table(doc, '{{CARGO_1}}')
    if t0 is not None and data.get('cargos'):
        # o template tem 4 linhas CARGO_n/PERIODO_n: preencher e zerar marcadores não usados
        cargos = data['cargos']
        for i in range(1, 5):
            cargo = cargos[i-1] if i-1 < len(cargos) else None
            for p in all_paragraphs(doc):
                for run in p.runs:
                    if '{{CARGO_%d}}' % i in run.text:
                        run.text = run.text.replace('{{CARGO_%d}}' % i, cargo[0] if cargo else '')
                    if '{{PERIODO_%d}}' % i in run.text:
                        run.text = run.text.replace('{{PERIODO_%d}}' % i, cargo[1] if cargo else '')

    # 5 tabelas de checklist (inserir nas âncoras)
    def evtab(key, anchor):
        rows, total, interp = pl['tabelas'][key]
        body = [[r[0], r[1], r[2]] for r in rows]
        body.append(['', 'TOTAL DE PONTOS:', str(total)])          # total (nº curto) na col Pontos
        body.append(['', 'INTERPRETAÇÃO: ' + str(interp), ''])     # texto longo na col larga
        # Item estreito · Descrição larga (justificada) · Pontos estreito
        if not insert_table_at_anchor(doc, anchor, ['Item', 'Descrição', 'Pontos'], body,
                                      widths=[0.7, 4.9, 0.8], justify_cols=(1,)):
            warnings.append('âncora %s não encontrada' % anchor)
        if not rows: warnings.append('aba %s vazia na planilha' % key)
    evtab('BM', '{{TABELA_AV_BIOMECANICA}}')
    evtab('MS', '{{TABELA_AV_MEMBROS_SUPERIORES}}')
    evtab('CV', '{{TABELA_AV_COLUNA_VERTEBRAL}}')
    for key, anchor, h1 in [('COND', '{{TABELA_COND_RUIM}}', 'Condição'),
                            ('ASECV', '{{TABELA_ASECV}}', 'Situação')]:
        rows = pl['tabelas'][key]
        # Condição/Situação larga (justificada) · SIM/NÃO estreitas (só "x")
        if not insert_table_at_anchor(doc, anchor, [h1, 'SIM', 'NÃO'], rows,
                                      widths=[5.0, 0.7, 0.7], justify_cols=(0,)):
            warnings.append('âncora %s não encontrada' % anchor)

    # validações
    full = '\n'.join(p.text for p in all_paragraphs(doc))
    residual = sorted(set(re.findall(r'\{\{[^}]+\}\}', full)))
    if residual: warnings.append('MARCADORES RESIDUAIS: ' + ', '.join(residual))
    perito = data.get('perito_nome', 'Irineu de Freitas Branco Junior')
    if perito not in full: warnings.append('IDENTIDADE: perito (%s) ausente' % perito)
    for bad in data.get('nomes_proibidos', []):
        if bad in full: warnings.append('VAZAMENTO: "%s" presente' % bad)

    doc.save(out_path)
    print('OK ->', out_path)
    print('Níveis: BM=%s · MS=%s · CV=%s → %s' %
          (pl['niveis']['BM'], pl['niveis']['MS'], pl['niveis']['CV'], pl['qualificacao']))
    if warnings:
        print('\n⚠ AVISOS:')
        for w in warnings: print('  -', w)
    else:
        print('Sem marcadores residuais. Identidade OK.')
    return not residual

if __name__ == '__main__':
    if len(sys.argv) != 5:
        print('uso: python3 build_laudo_ergo.py <template> <data.json> <planilha.xlsx> <saida.docx>'); sys.exit(1)
    build(sys.argv[1], sys.argv[2], sys.argv[3], sys.argv[4])
